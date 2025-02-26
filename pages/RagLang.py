import os
from io import BytesIO
import streamlit as st
import pandas as pd
from docx import Document as DocxDocument
import re  # Import the regular expression module

# Langchain
from langchain.embeddings import OllamaEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, TextLoader  # Import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import Ollama
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.prompts import PromptTemplate  # Import PromptTemplate
from langchain.chains.conversational_retrieval.prompts import CONDENSE_QUESTION_PROMPT
from langchain.chains import ConversationalRetrievalChain
from langchain.chains import LLMChain
from langchain.chains.combine_documents.stuff import create_stuff_documents_chain



# Configuration
OLLAMA_HOST = "http://10.1.11.60:11434"
LLM_MODEL = "deepseek-r1"
VECTOR_DB_PATH = "vector_db"


# --- Helper Functions ---
def load_document(file):
    """Loads a document based on its file type."""
    try:
        documents = []  # Initialize an empty list to hold documents

        if file.type == "application/pdf":
            loader = PyPDFLoader(file)
            loaded_documents = loader.load()
            for doc in loaded_documents:
                doc.metadata["file_name"] = file.name  # Add file name to metadata
            documents.extend(loaded_documents)
        elif file.type == "text/plain":
            loader = TextLoader(file)
            loaded_documents = loader.load()
            for doc in loaded_documents:
                doc.metadata["file_name"] = file.name  # Add file name to metadata
            documents.extend(loaded_documents)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # DOCX
            try:
                docx_doc = DocxDocument(file)
                text = "\n".join([para.text for para in docx_doc.paragraphs])
                doc = Document(page_content=text, metadata={"file_name": file.name})  # Include metadata
                documents.append(doc)
            except Exception as e:
                print(f"Error processing DOCX: {e}")
                st.error(f"Error processing DOCX: {e}") # show error on streamlit
                return []
        elif file.type == "text/csv":
            df = pd.read_csv(file)
            text = df.to_string()  # Convert DataFrame to string
            doc = Document(page_content=text, metadata={"file_name": file.name})
            documents.append(doc)

        elif file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            try:
                xls = pd.ExcelFile(file)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    text = f"Sheet: {sheet_name}\n{df.to_string()}"
                    doc = Document(page_content=text, metadata={"file_name": file.name, "sheet_name": sheet_name})
                    documents.append(doc)
            except Exception as e:
                print(f"Error processing XLSX: {e}")
                st.error(f"Error processing XLSX: {e}")
                return []
        else:
            raise ValueError(f"Unsupported file type: {file.type}")

        return documents

    except Exception as e:
        print(f"Error loading document: {e}")
        st.error(f"Error loading document: {e}") # show error on streamlit
        return []


def chunk_documents(documents, chunk_size=500, chunk_overlap=50):
    """Splits documents into smaller chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )
    return text_splitter.split_documents(documents)


def embed_and_store_chunks(chunks):
    """Embeds document chunks and stores them in a FAISS vector database."""
    embeddings = OllamaEmbeddings(model=LLM_MODEL, base_url=OLLAMA_HOST)
    #embeddings = HuggingFaceEmbeddings()

    vector_db = FAISS.from_documents(chunks, embeddings)
    vector_db.save_local(VECTOR_DB_PATH)
    return vector_db


def load_existing_db():
    """Loads an existing FAISS vector database if it exists."""
    try:
        embeddings = OllamaEmbeddings(model=LLM_MODEL, base_url=OLLAMA_HOST)
        vector_db = FAISS.load_local(VECTOR_DB_PATH, embeddings)
        return vector_db
    except Exception as e:
        print(f"Error loading existing database: {e}")
        return None

def remove_think_tags(text):
    """Removes <think>...</think> tags from the given text using regular expressions."""
    pattern = r'<think>.*?</think>'
    return re.sub(pattern, '', text, flags=re.DOTALL)

# --- Streamlit UI ---
st.set_page_config(page_title="📚 RAG Chatbot (Ollama)", layout="wide")
st.title("📚 Local RAG Chatbot using Ollama!")

# Sidebar for File Upload
st.sidebar.header("📤 Upload Documents")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDFs, TXT, DOCX, CSV, or XLSX",
    type=["pdf", "txt", "docx", "csv", "xlsx"],
    accept_multiple_files=True,
)

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Initialize conversation history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Document info
if "document_info" not in st.session_state:
    st.session_state.document_info = ""


# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])


# --- Main Logic ---
# Load existing DB at the beginning
vector_db = load_existing_db()
if vector_db:
    st.success("Existing database loaded.")

# File selection checkboxes
selected_files = {}
if uploaded_files:
    st.sidebar.subheader("Select files for context:")
    for file in uploaded_files:
        checkbox_key = file.name  # Use just the filename as the key

        # Initialize checkbox state if it doesn't exist
        if checkbox_key not in st.session_state:
            st.session_state[checkbox_key] = True

        selected = st.sidebar.checkbox(file.name, key=checkbox_key)


    files_to_process = [file for file in uploaded_files if st.session_state[file.name]]
    process_files = bool(files_to_process)  # Only process if there are selected files
else:
    files_to_process = []
    process_files = False

if process_files:
    with st.spinner("Processing documents..."):
        try:
            # Load documents
            documents = []
            for file in files_to_process:
                try:
                    loaded_docs = load_document(file)
                    if loaded_docs:  # Check if loaded_docs is not empty
                        if isinstance(loaded_docs, list):
                            documents.extend(loaded_docs)  # Extend if it's a list
                        else:
                            documents.append(loaded_docs)  # Append if it's a single Document object
                    else:
                        st.warning(f"No documents loaded from {file.name}")


                except Exception as e:
                    st.error(f"Error loading {file.name}: {e}")

            # Chunk documents
            if documents:  # Check if there are any documents to chunk
                chunks = chunk_documents(documents)

                # Embed and store chunks
                vector_db = embed_and_store_chunks(chunks)
                st.success("Documents processed and database created successfully!")

                # Store document info in session state
                document_names = ", ".join([file.name for file in files_to_process])
                st.session_state.document_info = f"You are chatting about the document(s): {document_names}"

                # Add the initial message to the chat history
                st.session_state.messages.append({"role": "assistant", "content": st.session_state.document_info})
                with st.chat_message("assistant"):
                    st.markdown(st.session_state.document_info)

            else:
                st.warning("No documents were loaded, cannot create the knowledge base.")
                vector_db = None  # Ensure vector_db is None if no documents are loaded

        except Exception as e:
            st.error(f"An error occurred during processing: {e}")


# Chat UI
if vector_db:
    retriever = vector_db.as_retriever(search_kwargs={"k": 3})

    # Custom Prompt (Optional)
    prompt_template = """{document_info}
        1. Use ONLY the context below.
        2. If unsure, say "I don’t know".
        3. Keep answers under 4 sentences.
        Context: {context}

        Question: {question}

        Answer:
    """
    PROMPT = PromptTemplate(
        template=prompt_template, input_variables=["context", "question", "document_info"]
    )
    chain_type_kwargs = {"prompt": PROMPT}

    # Set up the conversational chain
    llm = Ollama(model=LLM_MODEL, temperature=0, base_url=OLLAMA_HOST)
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever=retriever,
        condense_question_prompt=CONDENSE_QUESTION_PROMPT,
        return_source_documents=True,
        chain_type="stuff",
        combine_docs_chain_kwargs=chain_type_kwargs
    )
    # temp test ----------------
    # Create chain
    # llm_chain = LLMChain(llm=llm, prompt=PROMPT)
    # # Combine document chunks
    # document_chain = create_stuff_documents_chain(
    #     llm=llm,
    #     prompt=PROMPT
    # )
    # # temp
    # qa_temp = RetrievalQA.from_chain_type(
    #     llm=llm,
    #     retriever=retriever,
    #     chain_type="stuff", 
    # )
    # # temp test ----------------
    
    
    # user_input = st.text_input("Ask a question:")

    # if user_input:
    #     with st.spinner("Thinking..."):
    #         try:
    #             response = qa_temp.run(user_input)  
    #             st.write(response)
    #         except Exception as e:
    #             st.error(f"Error: {str(e)}")
                

    # Accept user input
    if query := st.chat_input("Ask me anything about the documents"):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": query})
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(query)

        # Get LLM response
        response = qa_chain({"question": query, "chat_history": st.session_state.chat_history, "document_info": st.session_state.document_info})
        # Remove <think>...</think> tags from the response
        response_content = remove_think_tags(response["answer"])

        # Add assistant response to chat history
        st.session_state.messages.append({"role": "assistant", "content": response_content})
        # Display assistant response in chat message container
        with st.chat_message("assistant"):
            st.markdown(response_content)

        # Update chat history
        st.session_state.chat_history.append((query, response_content))


else:
    st.info("Please upload documents to create the knowledge base.")