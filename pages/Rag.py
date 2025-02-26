import os
from io import BytesIO
import streamlit as st
import pandas as pd
from docx import Document as DocxDocument
import re  # Import the regular expression module

# Langchain
from langchain.embeddings import OllamaEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, TextLoader  # Import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import Ollama
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.prompts import PromptTemplate  # Import PromptTemplate


# Configuration
OLLAMA_HOST = "http://10.1.11.60:11434"
LLM_MODEL = "deepseek-r1"
VECTOR_DB_PATH = "vector_db"


# --- Helper Functions ---
def load_document(file):
    """Loads a document based on its file type."""
    try:
        documents = []  # Initialize an empty list to hold documents
        if file.type == "application/pdf":
            loader = PyPDFLoader(file)
            documents.extend(loader.load())  # Extend the documents list
        elif file.type == "text/plain":
            loader = TextLoader(file)
            documents.extend(loader.load())  # Extend the documents list
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # DOCX
            try:
                docx_doc = DocxDocument(file)
                text = "\n".join([para.text for para in docx_doc.paragraphs])
                #loader = TextLoader(BytesIO(text.encode('utf-8')), encoding='utf-8')  # Using TextLoader with BytesIO
                documents.append(Document(page_content=text))  # Create document directly
            except Exception as e:
                print(f"Error processing DOCX: {e}")
                st.error(f"Error processing DOCX: {e}") # show error on streamlit
                return []
        elif file.type == "text/csv":
            df = pd.read_csv(file)
            text = df.to_string()  # Convert DataFrame to string
            loader = TextLoader(BytesIO(text.encode('utf-8')), encoding='utf-8')  # Using TextLoader with BytesIO
            documents.extend(loader.load())  # Extend the documents list
        elif file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":  # XLSX
            try:
                xls = pd.ExcelFile(file)
                text = ""
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    text += f"Sheet: {sheet_name}\n{df.to_string()}\n"
                #loader = TextLoader(BytesIO(text.encode('utf-8')), encoding='utf-8')  # Using TextLoader with BytesIO
                documents.append(Document(page_content=text)) # Create document directly
            except Exception as e:
                print(f"Error processing XLSX: {e}")
                st.error(f"Error processing XLSX: {e}")  # show error on streamlit
                return []
        else:
            raise ValueError(f"Unsupported file type: {file.type}")
        return documents

    except Exception as e:
        print(f"Error loading document: {e}")
        st.error(f"Error loading document: {e}") # show error on streamlit
        return []


def chunk_documents(documents, chunk_size=500, chunk_overlap=50):
    """Splits documents into smaller chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )
    return text_splitter.split_documents(documents)


def embed_and_store_chunks(chunks):
    """Embeds document chunks and stores them in a FAISS vector database."""
    embeddings = OllamaEmbeddings(model=LLM_MODEL, base_url=OLLAMA_HOST)
    vector_db = FAISS.from_documents(chunks, embeddings)
    vector_db.save_local(VECTOR_DB_PATH)
    return vector_db


def load_existing_db():
    """Loads an existing FAISS vector database if it exists."""
    try:
        embeddings = OllamaEmbeddings(model=LLM_MODEL, base_url=OLLAMA_HOST)
        vector_db = FAISS.load_local(VECTOR_DB_PATH, embeddings)
        return vector_db
    except Exception as e:
        print(f"Error loading existing database: {e}")
        return None

def remove_think_tags(text):
    """Removes <think>...</think> tags from the given text using regular expressions."""
    pattern = r'<think>.*?</think>'
    return re.sub(pattern, '', text, flags=re.DOTALL)

# --- Streamlit UI ---
st.set_page_config(page_title="📚 RAG Chatbot (Ollama)", layout="wide")
st.title("📚 Local RAG Chatbot using Ollama!")

# Sidebar for File Upload
st.sidebar.header("📤 Upload Documents")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDFs, TXT, DOCX, CSV, or XLSX",
    type=["pdf", "txt", "docx", "csv", "xlsx"],
    accept_multiple_files=True,
)

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])


# --- Main Logic ---
vector_db = load_existing_db()  # Try to load existing DB first
process_files = False # use for decide to process files base on action.

if uploaded_files:
    process_files = True
    files_to_process = uploaded_files

if process_files:
    with st.spinner("Processing documents..."):
        try:
            # Load documents
            documents = []
            for file in files_to_process:
                try:
                    loaded_docs = load_document(file)
                    if loaded_docs:  # Check if loaded_docs is not empty
                        if isinstance(loaded_docs, list):
                            documents.extend(loaded_docs)  # Extend if it's a list
                        else:
                            documents.append(loaded_docs)  # Append if it's a single Document object
                    else:
                        st.warning(f"No documents loaded from {file.name}")


                except Exception as e:
                    st.error(f"Error loading {file.name}: {e}")

            # Chunk documents
            if documents:  # Check if there are any documents to chunk
                chunks = chunk_documents(documents)

                # Embed and store chunks
                vector_db = embed_and_store_chunks(chunks)
                st.success("Documents processed and database created successfully!")
            else:
                st.warning("No documents were loaded, cannot create the knowledge base.")
                vector_db = None  # Ensure vector_db is None if no documents are loaded

        except Exception as e:
            st.error(f"An error occurred during processing: {e}")


# Chat UI
if vector_db:
    retriever = vector_db.as_retriever()

    # Custom Prompt (Optional)
    prompt_template = """Use the following pieces of context to answer the question at the end. 
    If you don't know the answer, just say that you don't know, don't try to make up an answer.

    {context}

    Question: {question}
    """
    PROMPT = PromptTemplate(
        template=prompt_template, input_variables=["context", "question"]
    )
    chain_type_kwargs = {"prompt": PROMPT}


    # Accept user input
    if query := st.chat_input("Ask me anything about the documents"):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": query})
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(query)

        # Get LLM response
        llm = Ollama(model=LLM_MODEL, temperature=0, base_url=OLLAMA_HOST)
        qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever, chain_type_kwargs=chain_type_kwargs)
        response = qa_chain.run(query)
        # Remove <think>...</think> tags from the response
        response = remove_think_tags(response)

        # Add assistant response to chat history
        st.session_state.messages.append({"role": "assistant", "content": response})
        # Display assistant response in chat message container
        with st.chat_message("assistant"):
            st.markdown(response)

else:
    st.info("Please upload documents to create the knowledge base.")